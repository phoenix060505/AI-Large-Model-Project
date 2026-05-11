import os
import re
import json
import time
import hashlib
import requests

from io import BytesIO
from bs4 import BeautifulSoup
from tqdm import tqdm
from ddgs import DDGS
from openai import OpenAI

import fitz
import docx


class AICampusQACrawler:
    def __init__(
        self,
        user_requirement,
        allowed_domains=None,
        seed_urls=None,
        max_search_keywords=20,
        max_urls_per_keyword=15,
        max_pages=60,
        max_qa_per_chunk=8,
        output_file="sustech_qa_pairs.jsonl",
        raw_output_file="sustech_raw_documents.jsonl",
        model="deepseek-chat",
        delay=1.0,
        skip_existing_urls=True,
        skip_existing_questions=True,
        skip_duplicate_content=True
    ):
        self.user_requirement = user_requirement
        self.allowed_domains = allowed_domains or []
        self.seed_urls = seed_urls or []

        self.max_search_keywords = max_search_keywords
        self.max_urls_per_keyword = max_urls_per_keyword
        self.max_pages = max_pages
        self.max_qa_per_chunk = max_qa_per_chunk

        self.output_file = output_file
        self.raw_output_file = raw_output_file
        self.model = model
        self.delay = delay

        self.skip_existing_urls = skip_existing_urls
        self.skip_existing_questions = skip_existing_questions
        self.skip_duplicate_content = skip_duplicate_content

        self.client = OpenAI(
            api_key=os.getenv("DEEPSEEK_API_KEY"),
            base_url="https://api.deepseek.com"
        )

        self.headers = {
            "User-Agent": "Mozilla/5.0 (Educational AI Crawler Demo)"
        }

        # 单次运行 + 历史文件加载后的去重集合
        self.visited_urls = set()
        self.saved_questions = set()
        self.saved_doc_hashes = set()

    # =========================
    # 通用工具函数
    # =========================

    def clean_json_text(self, text):
        """
        清理模型返回内容，避免 ```json 包裹导致解析失败。
        """
        text = text.strip()
        text = re.sub(r"^```json", "", text).strip()
        text = re.sub(r"^```", "", text).strip()
        text = re.sub(r"```$", "", text).strip()
        return text

    def ask_ai_json(self, prompt):
        """
        调用 DeepSeek，并解析 JSON。
        """
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "你是一个严谨的数据处理助手。"
                            "你必须严格输出合法 JSON。"
                            "不要输出 Markdown，不要输出解释，不要使用 ``` 包裹。"
                        )
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.2
            )

            text = response.choices[0].message.content.strip()
            text = self.clean_json_text(text)
            return json.loads(text)

        except Exception as e:
            print(f"[AI JSON parse failed] {e}")
            return {}

    def save_jsonl(self, file_path, record):
        """
        追加保存 JSONL 文件。
        """
        with open(file_path, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")

    def make_id(self, text):
        """
        根据文本生成稳定 ID。
        """
        text = str(text)
        return hashlib.md5(text.encode("utf-8")).hexdigest()

    def normalize_url(self, url):
        """
        简单规范化 URL。
        """
        if not url:
            return ""

        url = url.strip()
        url = url.split("#")[0]
        url = url.rstrip("/")
        return url

    def normalize_question(self, question):
        """
        规范化问题文本，用于问题去重。
        """
        question = str(question).strip()
        question = re.sub(r"\s+", "", question)
        question = question.replace("？", "?")
        return question

    def get_content_hash(self, content):
        """
        根据正文内容生成 hash。
        只取前 8000 字符，避免超长文本导致处理过慢。
        """
        content = str(content).strip()
        content = re.sub(r"\s+", "", content)
        return self.make_id(content[:8000])

    def domain_allowed(self, url):
        """
        判断 URL 是否在允许域名内。
        如果 allowed_domains 为空，则不限制。
        """
        if not self.allowed_domains:
            return True

        return any(domain in url for domain in self.allowed_domains)

    # =========================
    # 读取历史文件，实现跨运行去重
    # =========================

    def load_existing_jsonl(self, file_path):
        """
        读取已有 JSONL 文件。
        """
        records = []

        if not os.path.exists(file_path):
            return records

        with open(file_path, "r", encoding="utf-8") as f:
            for line_idx, line in enumerate(f, start=1):
                line = line.strip()

                if not line:
                    continue

                try:
                    records.append(json.loads(line))
                except json.JSONDecodeError:
                    print(f"[Warning] JSONL 第 {line_idx} 行解析失败，已跳过：{file_path}")
                    continue

        return records

    def load_existing_state(self):
        """
        从已有 raw_documents 和 qa_pairs 文件中加载历史去重信息。
        """
        raw_records = self.load_existing_jsonl(self.raw_output_file)

        for record in raw_records:
            url = record.get("url", "").strip()
            if url:
                self.visited_urls.add(self.normalize_url(url))

            content = record.get("content", "").strip()
            if content:
                self.saved_doc_hashes.add(self.get_content_hash(content))

        qa_records = self.load_existing_jsonl(self.output_file)

        for record in qa_records:
            question = record.get("question", "").strip()
            if question:
                normalized_question = self.normalize_question(question)
                self.saved_questions.add(self.make_id(normalized_question))

        print(f"[Loaded existing state] 历史 raw documents: {len(raw_records)} 条")
        print(f"[Loaded existing state] 历史 QA pairs: {len(qa_records)} 条")
        print(f"[Loaded existing state] 已记录历史 URL: {len(self.visited_urls)} 个")
        print(f"[Loaded existing state] 已记录历史 content hash: {len(self.saved_doc_hashes)} 个")
        print(f"[Loaded existing state] 已记录历史 question hash: {len(self.saved_questions)} 个")

    # =========================
    # AI 生成搜索关键词
    # =========================

    def generate_search_keywords(self):
        prompt = f"""
用户想要构建一个南方科技大学校园知识库 QA 数据集。

用户需求：
{self.user_requirement}

请生成适合搜索网页的关键词。

要求：
1. 关键词要具体，不要太泛。
2. 每个主题至少生成 2-3 个不同关键词。
3. 优先覆盖：校园卡、宿舍、选课、加课、图书馆、食堂、校医院、学生事务、校园网、Blackboard、学生手册、新生攻略。
4. 关键词中尽量包含“南方科技大学”或“SUSTech”。
5. 可以包含“PDF”“DOCX”“学生手册”“办事指南”“管理细则”“通知”等词。
6. 最多生成 {self.max_search_keywords} 个关键词。
7. 只返回 JSON 对象，格式如下：

{{
  "keywords": [
    "南方科技大学 校园卡 挂失 补办",
    "南方科技大学 本科生选课办法 DOCX",
    "南方科技大学 学生手册 PDF"
  ]
}}
"""

        result = self.ask_ai_json(prompt)
        keywords = result.get("keywords", [])

        if not isinstance(keywords, list):
            print("[Warning] keywords 不是数组，使用默认关键词。")
            keywords = []

        default_keywords = [
            "南方科技大学 学生手册 PDF",
            "南方科技大学 新生攻略 PDF",
            "南方科技大学 校园卡 挂失 补办",
            "南方科技大学 本科生选课办法 DOCX",
            "南方科技大学 加课 选课 流程",
            "南方科技大学 图书馆 开放时间 借阅规则",
            "南方科技大学 学生事务 办事指南",
            "南方科技大学 Blackboard 校园网",
            "南方科技大学 学生公寓管理细则",
            "南方科技大学 校医院 医保 就诊",
            "南方科技大学 食堂 营业时间",
            "南方科技大学 学生事务中心 成绩单",
            "南方科技大学 校园网 使用指南",
            "南方科技大学 本科生管理规定",
            "南方科技大学 学籍管理规定"
        ]

        for kw in default_keywords:
            if kw not in keywords:
                keywords.append(kw)

        return keywords[: self.max_search_keywords]

    # =========================
    # 搜索网页
    # =========================

    def search_web(self, keyword):
        """
        使用 DDGS 搜索网页。
        """
        results = []

        try:
            with DDGS() as ddgs:
                for r in ddgs.text(
                    keyword,
                    max_results=self.max_urls_per_keyword
                ):
                    url = r.get("href", "")
                    if not url:
                        continue

                    url = self.normalize_url(url)

                    results.append({
                        "title": r.get("title", ""),
                        "url": url,
                        "snippet": r.get("body", "")
                    })

        except Exception as e:
            print(f"[Search failed] {keyword}: {e}")

        return results

    # =========================
    # AI 判断 URL 相关性
    # =========================

    def judge_url_relevance(self, title, snippet, url):
        prompt = f"""
请判断下面这个网页是否适合用于构建南方科技大学校园知识库 QA 数据集。

用户需求：
{self.user_requirement}

网页标题：
{title}

网页摘要：
{snippet}

网页 URL：
{url}

判断标准：
1. 和用户需求明显相关才返回 true。
2. 学校官网、学院官网、书院官网、图书馆官网、学生手册、通知、FAQ、DOCX、PDF 更优先。
3. 广告、商品页面、无关新闻、纯宣传页面、论坛灌水页面返回 false。
4. 如果是南方科技大学官方域名下的校园服务、学生事务、教学制度、图书馆、住宿、校园网等内容，通常应返回 true。

请只返回 JSON 对象：
{{
  "relevant": true,
  "reason": "简短原因",
  "category": "library"
}}
"""

        result = self.ask_ai_json(prompt)

        if "relevant" not in result:
            result["relevant"] = False

        return result

    # =========================
    # 下载并解析 HTML / PDF / DOCX
    # =========================

    def fetch_content(self, url):
        """
        下载网页 / PDF / DOCX，并返回 title, content。
        """
        try:
            resp = requests.get(
                url,
                headers=self.headers,
                timeout=30
            )

            resp.raise_for_status()

            content_type = resp.headers.get("Content-Type", "").lower()
            url_lower = url.lower()

            # HTML
            if (
                "text/html" in content_type
                or url_lower.endswith(".html")
                or url_lower.endswith(".htm")
                or "." not in url_lower.split("/")[-1]
            ):
                resp.encoding = resp.apparent_encoding
                title, content = self.extract_html_text(resp.text)
                return title, content

            # PDF
            if "application/pdf" in content_type or url_lower.endswith(".pdf"):
                return self.extract_pdf_text(resp.content, url)

            # DOCX
            if (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type
                or url_lower.endswith(".docx")
            ):
                return self.extract_docx_text(resp.content, url)

            # 老式 .doc 暂不处理
            if url_lower.endswith(".doc"):
                print(f"[Skip DOC] 暂不支持 .doc 文件：{url}")
                return None, None

            print(f"[Skip unsupported content type] {url} | {content_type}")
            return None, None

        except Exception as e:
            print(f"[Fetch failed] {url}: {e}")
            return None, None

    def extract_html_text(self, html):
        """
        从 HTML 中提取正文。
        """
        soup = BeautifulSoup(html, "lxml")

        for tag in soup([
            "script",
            "style",
            "nav",
            "footer",
            "header",
            "aside",
            "form",
            "noscript"
        ]):
            tag.decompose()

        title = soup.title.get_text(strip=True) if soup.title else ""

        main = (
            soup.find("main")
            or soup.find("article")
            or soup.find("div", class_=re.compile("content|main|article", re.I))
            or soup.body
            or soup
        )

        text = main.get_text("\n", strip=True)

        lines = []
        for line in text.split("\n"):
            line = re.sub(r"\s+", " ", line).strip()
            if len(line) >= 8:
                lines.append(line)

        content = "\n".join(lines)

        return title, content

    def extract_pdf_text(self, binary_content, url):
        """
        提取 PDF 文本。
        """
        try:
            pdf = fitz.open(stream=binary_content, filetype="pdf")
            texts = []

            for page_idx, page in enumerate(pdf):
                page_text = page.get_text("text")
                page_text = page_text.strip()

                if page_text:
                    texts.append(f"第 {page_idx + 1} 页\n{page_text}")

            title = url.split("/")[-1] or "PDF Document"
            content = "\n\n".join(texts)

            return title, content

        except Exception as e:
            print(f"[PDF parse failed] {url}: {e}")
            return None, None

    def extract_docx_text(self, binary_content, url):
        """
        提取 DOCX 文本。
        """
        try:
            document = docx.Document(BytesIO(binary_content))
            paragraphs = []

            for p in document.paragraphs:
                text = p.text.strip()
                if text:
                    paragraphs.append(text)

            # 尝试读取表格内容
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))

            title = url.split("/")[-1] or "DOCX Document"
            content = "\n".join(paragraphs)

            return title, content

        except Exception as e:
            print(f"[DOCX parse failed] {url}: {e}")
            return None, None

    # =========================
    # 文本切块
    # =========================

    def split_text(self, text, chunk_size=2500, overlap=300):
        """
        把长文本切成多个 chunk。
        """
        chunks = []
        text = text.strip()

        if not text:
            return chunks

        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(start + chunk_size, text_len)
            chunk = text[start:end].strip()

            if len(chunk) >= 300:
                chunks.append(chunk)

            if end >= text_len:
                break

            start = end - overlap
            if start < 0:
                start = 0

        return chunks

    # =========================
    # AI 生成 QA 对
    # =========================

    def generate_qa_pairs(self, title, url, content):
        prompt = f"""
请根据下面的南方科技大学校园资料生成 QA 对。

用户需求：
{self.user_requirement}

资料标题：
{title}

资料 URL：
{url}

资料正文：
{content}

生成要求：
1. 只能根据资料正文生成，不允许编造。
2. 如果正文中没有明确答案，不要生成对应问题。
3. 问题要自然，像学生真实会问的问题。
4. 答案要准确、简洁。
5. 尽量覆盖不同信息点，不要生成重复问题。
6. 每条 QA 都要带 source_title 和 source_url。
7. 如果信息具有时效性，请在答案末尾补充“具体安排以学校最新官方通知为准”。
8. 最多生成 {self.max_qa_per_chunk} 条 QA。
9. 如果该 chunk 没有有价值的信息，可以返回空数组。

请只返回 JSON 对象：
{{
  "qa_pairs": [
    {{
      "question": "问题",
      "answer": "答案",
      "category": "分类，例如 course / dormitory / library / campus_card / network / medical / student_affairs / other",
      "source_title": "资料标题",
      "source_url": "资料URL"
    }}
  ]
}}
"""

        result = self.ask_ai_json(prompt)
        qa_pairs = result.get("qa_pairs", [])

        if not isinstance(qa_pairs, list):
            return []

        cleaned = []

        for qa in qa_pairs:
            if not isinstance(qa, dict):
                continue

            question = str(qa.get("question", "")).strip()
            answer = str(qa.get("answer", "")).strip()

            if len(question) < 5 or len(answer) < 5:
                continue

            qa["question"] = question
            qa["answer"] = answer
            qa["source_title"] = qa.get("source_title", title)
            qa["source_url"] = qa.get("source_url", url)
            qa["category"] = qa.get("category", "other")

            cleaned.append(qa)

        return cleaned

    # =========================
    # 主流程
    # =========================

    def run(self):
        # 关键新增：先读取历史文件，避免重复爬取和重复写入
        self.load_existing_state()

        print("\nStep 1: DeepSeek 正在根据需求生成搜索关键词...")
        keywords = self.generate_search_keywords()

        print("\n生成的搜索关键词：")
        for kw in keywords:
            print("-", kw)

        candidate_urls = []

        print("\nStep 2: 正在搜索相关网页...")
        for keyword in keywords:
            print(f"\n搜索关键词：{keyword}")

            search_results = self.search_web(keyword)

            for item in search_results:
                url = item["url"]

                if not url:
                    continue

                url = self.normalize_url(url)

                if self.skip_existing_urls and url in self.visited_urls:
                    print(f"  -> 跳过历史 URL：{url}")
                    continue

                if not self.domain_allowed(url):
                    continue

                print(f"DeepSeek 判断网页相关性：{item['title'][:60]}")

                judge = self.judge_url_relevance(
                    title=item["title"],
                    snippet=item["snippet"],
                    url=url
                )

                if judge.get("relevant") is True:
                    item["url"] = url
                    item["ai_reason"] = judge.get("reason", "")
                    item["ai_category"] = judge.get("category", "other")
                    candidate_urls.append(item)
                    print("  -> 保留")
                else:
                    print("  -> 跳过")

                time.sleep(0.3)

        # 加入人工指定 seed urls
        for url in self.seed_urls:
            url = self.normalize_url(url)

            if not self.domain_allowed(url):
                continue

            if self.skip_existing_urls and url in self.visited_urls:
                print(f"  -> 跳过历史 seed URL：{url}")
                continue

            candidate_urls.append({
                "title": "Seed URL",
                "url": url,
                "snippet": "",
                "ai_reason": "人工指定入口",
                "ai_category": "seed"
            })

        # 候选 URL 去重
        unique_candidates = []
        seen_urls = set()

        for item in candidate_urls:
            url = self.normalize_url(item["url"])
            if url and url not in seen_urls:
                item["url"] = url
                unique_candidates.append(item)
                seen_urls.add(url)

        candidate_urls = unique_candidates

        print(f"\nStep 3: 筛选后得到 {len(candidate_urls)} 个新的候选网页 / 文件。")

        crawled_count = 0
        skipped_duplicate_content_count = 0
        total_qa_count = 0

        for item in tqdm(candidate_urls, desc="Crawling and generating QA"):
            if crawled_count >= self.max_pages:
                break

            url = self.normalize_url(item["url"])

            if self.skip_existing_urls and url in self.visited_urls:
                continue

            # 先加入，避免本次运行重复处理同一 URL
            self.visited_urls.add(url)

            title, content = self.fetch_content(url)

            if not content or len(content) < 200:
                continue

            content_hash = self.get_content_hash(content)

            if self.skip_duplicate_content and content_hash in self.saved_doc_hashes:
                print(f"[Skip duplicate content] {url}")
                skipped_duplicate_content_count += 1
                continue

            self.saved_doc_hashes.add(content_hash)

            raw_record = {
                "doc_id": self.make_id(url),
                "url": url,
                "title": title,
                "content": content,
                "content_hash": content_hash,
                "content_length": len(content),
                "ai_category": item.get("ai_category", "other"),
                "ai_reason": item.get("ai_reason", ""),
                "crawl_time": time.strftime("%Y-%m-%d %H:%M:%S")
            }

            self.save_jsonl(self.raw_output_file, raw_record)

            chunks = self.split_text(content, chunk_size=2500, overlap=300)

            if not chunks:
                continue

            for chunk_id, chunk in enumerate(chunks):
                qa_pairs = self.generate_qa_pairs(
                    title=f"{title} - chunk {chunk_id + 1}",
                    url=url,
                    content=chunk
                )

                for qa in qa_pairs:
                    question = qa.get("question", "").strip()

                    if not question:
                        continue

                    normalized_question = self.normalize_question(question)
                    question_id = self.make_id(normalized_question)

                    if self.skip_existing_questions and question_id in self.saved_questions:
                        continue

                    self.saved_questions.add(question_id)

                    qa["id"] = question_id
                    qa["crawl_source"] = "deepseek_ai_auto_crawler"
                    qa["created_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
                    qa["chunk_id"] = chunk_id + 1
                    qa["source_doc_id"] = self.make_id(url)
                    qa["source_url"] = qa.get("source_url", url)

                    self.save_jsonl(self.output_file, qa)
                    total_qa_count += 1

                time.sleep(0.5)

            crawled_count += 1
            time.sleep(self.delay)

        print("\n完成。")
        print(f"新处理网页 / 文件数量：{crawled_count}")
        print(f"跳过重复内容数量：{skipped_duplicate_content_count}")
        print(f"新增 QA 对数量：{total_qa_count}")
        print(f"原始网页数据保存到：{self.raw_output_file}")
        print(f"QA 对保存到：{self.output_file}")


if __name__ == "__main__":
    requirement = """
我正在做南方科技大学校园知识库项目。
请帮我自动寻找适合构建 QA 对的数据网页。

重点关注：
1. 校园卡
2. 宿舍
3. 选课和加课
4. 图书馆
5. 食堂
6. 校医院
7. 学生事务
8. Blackboard 和校园网
9. 学生手册
10. 新生攻略
11. 办事指南
12. 教学管理规定
13. 学籍管理
14. 成绩单
15. 体育课和体测
16. 校园服务

要求：
1. 优先使用南方科技大学官网、学生手册、书院官网、图书馆官网、教学工作部、学生事务、官方通知等公开信息。
2. 生成的 QA 对要适合用于 RAG 校园问答系统。
3. 答案必须基于网页或文件原文，不允许编造。
"""

    crawler = AICampusQACrawler(
        user_requirement=requirement,
        allowed_domains=[
            "sustech.edu.cn",
            "www.sustech.edu.cn",
            "osa.sustech.edu.cn",
            "lib.sustech.edu.cn",
            "teaching.sustech.edu.cn",
            "gao.sustech.edu.cn",
            "student.sustech.edu.cn",
            "mirrors.sustech.edu.cn"
        ],
        seed_urls=[
            "https://www.sustech.edu.cn/zh/students.html",
            "https://lib.sustech.edu.cn/",
            "https://osa.sustech.edu.cn/",
            "https://gao.sustech.edu.cn/"
        ],
        max_search_keywords=30,
        max_urls_per_keyword=20,
        max_pages=1000,
        max_qa_per_chunk=10,
        output_file="sustech_qa_pairs.jsonl",
        raw_output_file="sustech_raw_documents.jsonl",
        model="deepseek-chat",
        delay=1.0,

        # 这三个参数控制历史去重
        skip_existing_urls=True,
        skip_existing_questions=True,
        skip_duplicate_content=True
    )

    crawler.run()